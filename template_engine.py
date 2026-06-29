import io
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_html_preview(data: dict, task_type: str) -> str:
    # Google Font for Thai (Sarabun)
    font_link = '<link href="https://fonts.googleapis.com/css2?family=Sarabun:wght@400;700&display=swap" rel="stylesheet">'
    
    html = f"""
    <html>
    <head>
        {font_link}
        <style>
            .a4-preview {{
                background: white;
                width: 210mm;
                min-height: 297mm;
                margin: 0 auto;
                padding: 20mm 20mm 20mm 30mm;
                box-shadow: 0 0 10px rgba(0,0,0,0.1);
                font-family: 'Sarabun', sans-serif;
                font-size: 16pt;
                color: black;
                line-height: 1.5;
            }}
            .center {{ text-align: center; }}
            .right {{ text-align: right; }}
            .indent {{ text-indent: 2.5cm; }}
            .bold {{ font-weight: bold; }}
            .header-img {{ width: 1.5in; display: block; margin: 0 auto; }}
        </style>
    </head>
    <body>
        <div class="a4-preview">
    """
    
    if task_type in ["Official Memos (บันทึกข้อความ)", "Letters to Parents (จดหมายแจ้งผู้ปกครอง)"]:
        title_text = "บันทึกข้อความ" if "Memos" in task_type else "จดหมายแจ้งผู้ปกครอง"
        html += f"""
        <h2 class="center bold">{title_text}</h2>
        <div><span class="bold">ส่วนราชการ</span> {data.get('government_agency', '................................')}</div>
        <div><span class="bold">ที่</span> {data.get('document_number', '................')} <span style="float:right;"><span class="bold">วันที่</span> {data.get('date', '................')}</span></div>
        <div><span class="bold">เรื่อง</span> {data.get('subject', '................')}</div>
        <hr style="border-top: 1px solid black; margin: 10px 0;">
        <div><span class="bold">เรียน</span> {data.get('to', '................')}</div>
        <br>
        """
        for paragraph in data.get('body', '').split('\n'):
            if paragraph.strip():
                html += f'<div class="indent">{paragraph}</div>'
        
        html += f"""
        <br><br><br>
        <div style="width: 50%; float: right; text-align: center;">
            (ลงชื่อ)......................................................<br>
            ({data.get('signature_name', '................................')})<br>
            {data.get('signature_title', '................................')}
        </div>
        <div style="clear: both;"></div>
        """
        
        # Render any custom sections the AI added
        core_keys = ['government_agency', 'document_number', 'date', 'subject', 'to', 'body', 'signature_name', 'signature_title']
        for key, value in data.items():
            if key not in core_keys:
                html += f"<br><br><div class='bold'>{key.replace('_', ' ').title()}</div>"
                for paragraph in str(value).split('\n'):
                    if paragraph.strip():
                        html += f'<div>{paragraph}</div>'
    else:
        # Generic fallback
        html += f"<h1 class='center'>{task_type}</h1>"
        for key, value in data.items():
            html += f"<b>{key.capitalize()}:</b><br>"
            for paragraph in str(value).split('\n'):
                if paragraph.strip():
                    html += f"<div class='indent'>{paragraph}</div>"
            html += "<br>"
            
    html += """
        </div>
    </body>
    </html>
    """
    return html

def generate_docx(data: dict, task_type: str) -> io.BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun PSK'
    font.size = Pt(16)
    
    if task_type in ["Official Memos (บันทึกข้อความ)", "Letters to Parents (จดหมายแจ้งผู้ปกครอง)"]:
        title_text = "บันทึกข้อความ" if "Memos" in task_type else "จดหมายแจ้งผู้ปกครอง"
        p = doc.add_paragraph(title_text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(29)
        
        doc.add_paragraph(f"ส่วนราชการ {data.get('government_agency', '')}")
        doc.add_paragraph(f"ที่ {data.get('document_number', '')} \t\t วันที่ {data.get('date', '')}")
        doc.add_paragraph(f"เรื่อง {data.get('subject', '')}")
        doc.add_paragraph(f"เรียน {data.get('to', '')}")
        
        for paragraph in data.get('body', '').split('\n'):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph)
                p.paragraph_format.first_line_indent = Inches(1.0)
                
        doc.add_paragraph()
        doc.add_paragraph()
        sign_p = doc.add_paragraph(f"(ลงชื่อ)......................................................")
        sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        name_p = doc.add_paragraph(f"({data.get('signature_name', '')})")
        name_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        title_p = doc.add_paragraph(f"{data.get('signature_title', '')}")
        title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Render any custom sections the AI added
        core_keys = ['government_agency', 'document_number', 'date', 'subject', 'to', 'body', 'signature_name', 'signature_title']
        for key, value in data.items():
            if key not in core_keys:
                doc.add_paragraph()
                p = doc.add_paragraph()
                run = p.add_run(key.replace('_', ' ').title())
                run.bold = True
                for paragraph in str(value).split('\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
    else:
        doc.add_heading(task_type, level=1)
        for key, value in data.items():
            doc.add_heading(key.capitalize(), level=2)
            doc.add_paragraph(str(value))
            
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
