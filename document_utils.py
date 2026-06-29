from docx import Document
import io

def create_word_document(content_markdown: str):
    """
    Creates a Word document from a simple markdown string.
    Returns a BytesIO object containing the docx file.
    """
    doc = Document()
    
    # Simple parser for markdown to handle basic formatting
    lines = content_markdown.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Check for ordered lists like "1. ", "2. ", etc.
        elif len(line) > 2 and line[0].isdigit() and line[1:3] == '. ':
            doc.add_paragraph(line[3:], style='List Number')
        else:
            if line.strip() != "":
                doc.add_paragraph(line)
            
    # Save to BytesIO stream
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
